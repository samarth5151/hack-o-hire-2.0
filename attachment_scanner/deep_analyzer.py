# attachment_scanner/deep_analyzer.py
"""
Deep content analysis for attachments using llama3:latest.

Supported file types:
  - PDF        → extract text via pdfminer/PyMuPDF, analyze with llama3:latest
  - Word (DOCX)→ extract text via python-docx, analyze with llama3:latest
  - Text files → pass content directly to llama3:latest
  - Audio      → call voice-scanner container API
  - Images     → static analysis (metadata + basic checks)
"""
from __future__ import annotations

import io
import os
import re
import json
import requests
from pathlib import Path
from typing import Dict, Any, List, Optional

_OLLAMA_MODEL = "llama3"
_OLLAMA_URL   = os.environ.get("OLLAMA_HOST", os.environ.get("OLLAMA_URL", "http://localhost:11434"))
_VOICE_API    = os.environ.get("VOICE_SCANNER_URL", "http://localhost:8006")

# Supported file extensions by category
PDF_EXTS    = {".pdf"}
WORD_EXTS   = {".doc", ".docx", ".docm"}
TEXT_EXTS   = {".txt", ".csv", ".log", ".xml", ".json", ".html", ".htm", ".eml", ".msg"}
AUDIO_EXTS  = {".wav", ".mp3", ".flac", ".ogg", ".m4a", ".aac", ".wma", ".mp4", ".webm"}
IMAGE_EXTS  = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"}

# ── Text extraction helpers ─────────────────────────────────────────────────

def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fast) or pdfminer fallback."""
    text = ""

    # Try PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = []
        for page in doc:
            parts.append(page.get_text())
        text = "\n".join(parts).strip()
        doc.close()
        if text:
            return text
    except Exception:
        pass

    # Fallback: pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        from pdfminer.layout import LAParams
        buf = io.BytesIO(file_bytes)
        text = pdfminer_extract(buf, laparams=LAParams())
        return text.strip()
    except Exception:
        pass

    return "[PDF text extraction failed]"


def _extract_word_text(file_bytes: bytes) -> str:
    """Extract text from DOCX/DOC file."""
    try:
        import docx
        buf = io.BytesIO(file_bytes)
        doc = docx.Document(buf)
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception:
        return "[Word text extraction failed]"


def _extract_text_content(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from text-based files."""
    try:
        # Try UTF-8 first, then latin-1
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1", errors="replace")
    except Exception:
        return "[Text extraction failed]"


# ── LLM Analysis ──────────────────────────────────────────────────────────────

_DEEP_ANALYSIS_PROMPT = """You are a cybersecurity analyst performing deep phishing analysis on a document extracted from an email attachment.

FILENAME: {filename}
FILE TYPE: {file_type}
EXTRACTED CONTENT (first 4000 chars):
{content}

Analyze this content for:
1. Phishing indicators (fake login pages, spoofed brands, urgency tactics)
2. Credentials (passwords, API keys, account numbers, credit cards)
3. Suspicious links (especially shortened URLs, lookalike domains)
4. Social engineering patterns

Return ONLY valid JSON (no markdown, no explanation):
{{
  "phishing_score": <0-100 integer>,
  "verdict": "CLEAN" or "SUSPICIOUS" or "PHISHING" or "MALWARE",
  "credentials_found": [
    {{"type": "password/api_key/credit_card/account_number/other", "value": "<redacted-if-sensitive>", "context": "surrounding text (max 50 chars)"}}
  ],
  "links_found": [
    {{"url": "full URL", "suspicious": true/false, "reason": "why suspicious or safe"}}
  ],
  "sensitive_data": ["SSN pattern found", "Credit card number detected", etc.],
  "threats_detected": ["Fake Microsoft login page", "Urgency tactics used", etc.],
  "impersonated_entity": "Brand being impersonated (or null)",
  "summary": "2-3 sentences describing the security risk of this document"
}}"""

_RULE_BASED_PROMPT = """Analyze this document content and identify which of these threat categories apply.

CONTENT:
{content}

Categories to check (return which ones apply):
1. Credential harvesting (asks for passwords, pins, credit cards)
2. Phishing link (contains links to fake login pages)
3. Social engineering (urgency, fear, authority tactics)
4. Brand impersonation (pretends to be PayPal, Microsoft, bank, etc.)
5. Financial fraud (wire transfers, gift cards, bitcoin requests)
6. Malware indicators (scripts, executable links, macro instructions)

Return ONLY valid JSON:
{{
  "rule_score": <0-100 integer>,
  "triggered_categories": [
    {{"category": "Credential harvesting", "evidence": "asks for password in form on page 2", "severity": "High"}}
  ],
  "safe_indicators": ["Legitimate company header", "Valid contact info"],
  "overall_assessment": "1 sentence summary"
}}"""


def _call_llama_for_content(prompt: str) -> Dict:
    """Call llama3:latest via ollama and parse JSON response."""
    try:
        import ollama
        resp = ollama.chat(
            model=_OLLAMA_MODEL,
            messages=[{"role": "user", "content": prompt}],
            options={"temperature": 0.1, "num_predict": 1024},
        )
        raw = resp["message"]["content"].strip()
        raw = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()
        if "```" in raw:
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw)
        start = raw.find("{")
        end   = raw.rfind("}") + 1
        if start >= 0 and end > start:
            raw = raw[start:end]
        return json.loads(raw)
    except Exception as exc:
        return {"error": str(exc), "llm_available": False}


# ── Voice analysis via API ─────────────────────────────────────────────────────

def _analyze_audio_via_api(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Send audio file to voice-scanner container API and return result."""
    try:
        files   = {"file": (filename, io.BytesIO(file_bytes), "audio/mpeg")}
        resp    = requests.post(
            f"{_VOICE_API}/analyze/voice",
            files=files,
            timeout=60,
        )
        if resp.status_code == 200:
            data = resp.json()
            return {
                "available":    True,
                "verdict":      data.get("verdict", "UNKNOWN"),
                "risk_score":   data.get("risk_score", 0),
                "risk_tier":    data.get("risk_tier", "LOW"),
                "confidence":   data.get("confidence", "0%"),
                "is_deepfake":  "FAKE" in str(data.get("verdict", "")).upper(),
                "analysis":     data,
            }
        else:
            return {
                "available": False,
                "error":     f"Voice API returned {resp.status_code}",
                "verdict":   "UNKNOWN",
                "risk_score": 0,
            }
    except Exception as exc:
        return {
            "available": False,
            "error":     str(exc),
            "verdict":   "UNKNOWN",
            "risk_score": 0,
        }


# ── Image analysis ─────────────────────────────────────────────────────────────

def _analyze_image(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Static analysis for image files (metadata + anomaly checks)."""
    findings = []
    score    = 0

    try:
        from PIL import Image
        import io as _io
        img = Image.open(_io.BytesIO(file_bytes))

        # Check for embedded EXIF data
        try:
            exif = img._getexif() or {}
            if exif:
                findings.append({"type": "EXIF metadata present", "detail": f"{len(exif)} EXIF tag(s) found", "severity": "Low"})
        except Exception:
            pass

        # Check dimensions (very large images can carry steganography)
        w, h = img.size
        if w * h > 10_000_000:
            findings.append({"type": "Large image dimensions", "detail": f"{w}x{h} — possible steganography payload", "severity": "Low"})
            score += 5

        # Check for suspicious format (JPEG with unexpected header)
        if filename.lower().endswith(".jpg") and not file_bytes.startswith(b"\xff\xd8"):
            findings.append({"type": "JPEG header mismatch", "detail": "File claims to be JPEG but magic bytes differ", "severity": "High"})
            score += 30

    except Exception:
        findings.append({"type": "Image parse error", "detail": "Could not decode image", "severity": "Info"})

    return {
        "type":         "image_static",
        "findings":     findings,
        "score":        score,
        "verdict":      "SUSPICIOUS" if score >= 20 else "CLEAN",
        "summary":      f"Image static analysis: {len(findings)} indicator(s) found.",
    }


# ── Main deep analysis entry point ─────────────────────────────────────────────

def analyze_deep(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Perform deep content analysis on an attachment.

    Returns structured result with:
      - extracted_text (first 500 chars preview)
      - llm_analysis (phishing score, credentials, links, threats)
      - rule_based (rule-based threat categories)
      - voice_analysis (for audio files)
      - file_type, overall_score, verdict, summary
    """
    ext      = ("." + filename.lower().rsplit(".", 1)[-1]) if "." in filename else ""
    result: Dict[str, Any] = {
        "filename":       filename,
        "file_extension": ext,
        "analysis_type":  "deep",
    }

    # ── Audio files → Voice Scanner API ──────────────────────────────────────
    if ext in AUDIO_EXTS:
        result["file_type"]      = "audio"
        result["voice_analysis"] = _analyze_audio_via_api(file_bytes, filename)
        va = result["voice_analysis"]
        result["overall_score"]  = va.get("risk_score", 0)
        result["verdict"]        = "FAKE_VOICE" if va.get("is_deepfake") else ("UNKNOWN" if not va.get("available") else "REAL_VOICE")
        result["summary"]        = (
            f"Voice deepfake analysis: {va.get('verdict', 'N/A')} "
            f"(risk score: {va.get('risk_score', 0)}/100)"
        )
        return result

    # ── Image files → Static analysis ────────────────────────────────────────
    if ext in IMAGE_EXTS:
        result["file_type"]      = "image"
        img_result               = _analyze_image(file_bytes, filename)
        result["image_analysis"] = img_result
        result["overall_score"]  = img_result["score"]
        result["verdict"]        = img_result["verdict"]
        result["summary"]        = img_result["summary"]
        return result

    # ── Text/Document files → extract content + LLM analysis ─────────────────
    extracted_text = ""

    if ext in PDF_EXTS:
        result["file_type"] = "pdf"
        extracted_text = _extract_pdf_text(file_bytes)
    elif ext in WORD_EXTS:
        result["file_type"] = "word"
        extracted_text = _extract_word_text(file_bytes)
    elif ext in TEXT_EXTS:
        result["file_type"] = "text"
        extracted_text = _extract_text_content(file_bytes, filename)
    else:
        # Try as plain text for unknown types
        result["file_type"] = "unknown"
        try:
            extracted_text = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            extracted_text = ""

    result["extracted_text_preview"] = extracted_text[:500] if extracted_text else ""
    result["extracted_text_length"]  = len(extracted_text)

    if not extracted_text or extracted_text.startswith("["):
        result["overall_score"] = 0
        result["verdict"]       = "UNKNOWN"
        result["summary"]       = "Could not extract text content from this file for deep analysis."
        return result

    # ── LLM deep analysis ─────────────────────────────────────────────────────
    llm_prompt = _DEEP_ANALYSIS_PROMPT.format(
        filename  = filename,
        file_type = result["file_type"],
        content   = extracted_text[:4000],
    )
    llm_data = _call_llama_for_content(llm_prompt)

    result["llm_analysis"] = {
        "phishing_score":    int(llm_data.get("phishing_score", 0)),
        "verdict":           llm_data.get("verdict", "UNKNOWN"),
        "credentials_found": llm_data.get("credentials_found", []),
        "links_found":       llm_data.get("links_found", []),
        "sensitive_data":    llm_data.get("sensitive_data", []),
        "threats_detected":  llm_data.get("threats_detected", []),
        "impersonated_entity": llm_data.get("impersonated_entity"),
        "summary":           llm_data.get("summary", ""),
        "model":             "llama3:latest",
        "available":         "error" not in llm_data,
    }

    # ── Rule-based analysis ───────────────────────────────────────────────────
    rule_prompt = _RULE_BASED_PROMPT.format(content=extracted_text[:3000])
    rule_data   = _call_llama_for_content(rule_prompt)

    result["rule_based"] = {
        "score":               int(rule_data.get("rule_score", 0)),
        "triggered_categories": rule_data.get("triggered_categories", []),
        "safe_indicators":      rule_data.get("safe_indicators", []),
        "assessment":           rule_data.get("overall_assessment", ""),
        "available":            "error" not in rule_data,
    }

    # ── Compute overall score ─────────────────────────────────────────────────
    llm_score  = float(result["llm_analysis"]["phishing_score"])
    rule_score = float(result["rule_based"]["score"])
    overall    = round(llm_score * 0.65 + rule_score * 0.35, 1)

    result["overall_score"] = overall
    result["verdict"]       = llm_data.get("verdict", "CLEAN")
    result["summary"]       = (
        result["llm_analysis"].get("summary") or
        f"Deep analysis complete. Phishing risk score: {overall}/100"
    )

    return result
