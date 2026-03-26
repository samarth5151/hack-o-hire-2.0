import pymupdf
import pytesseract
import docx as python_docx
from PIL import Image
import email
from email import policy
import io
import re
import zipfile
import time

pytesseract.pytesseract.tesseract_cmd = \
    r'C:\Program Files\Tesseract-OCR\tesseract.exe'


def extract_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    start = time.perf_counter()
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    total_pages = len(doc)
    print(f"[extract] PDF opened with {total_pages} page(s)")
    for i in range(total_pages):
        page_start = time.perf_counter()
        text_parts.append(f"[Page {i+1}]\n{doc[i].get_text()}")
        if total_pages <= 5 or (i + 1) == total_pages or (i + 1) % 5 == 0:
            print(
                f"[extract] Processed PDF page {i+1}/{total_pages} "
                f"in {time.perf_counter() - page_start:.2f}s"
            )
    print(f"[extract] PDF extraction finished in {time.perf_counter() - start:.2f}s")
    return "\n".join(text_parts)


def extract_from_image(file_bytes: bytes) -> str:
    image = Image.open(io.BytesIO(file_bytes))
    if image.mode != "RGB":
        image = image.convert("RGB")
    return pytesseract.image_to_string(image, lang="eng")


def extract_from_docx(file_bytes: bytes) -> str:
    doc = python_docx.Document(io.BytesIO(file_bytes))
    parts = []

    # paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # table text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    # embedded images — OCR each one
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            for name in z.namelist():
                if name.startswith("word/media/") and \
                   any(name.lower().endswith(e)
                       for e in (".png", ".jpg", ".jpeg", ".bmp", ".tiff")):
                    try:
                        img_bytes = z.read(name)
                        image = Image.open(io.BytesIO(img_bytes))
                        if image.mode != "RGB":
                            image = image.convert("RGB")
                        ocr = pytesseract.image_to_string(image, lang="eng")
                        if ocr.strip():
                            parts.append(
                                f"[Embedded image: {name}]\n{ocr}")
                    except Exception:
                        pass
    except zipfile.BadZipFile:
        pass

    return "\n".join(parts)


def extract_from_email(file_bytes: bytes) -> str:
    raw = file_bytes.decode("utf-8", errors="ignore")
    msg = email.message_from_string(raw, policy=policy.default)
    parts = [
        f"FROM: {msg.get('from', '')}",
        f"TO: {msg.get('to', '')}",
        f"SUBJECT: {msg.get('subject', '')}",
        "---BODY---",
    ]
    for part in msg.walk():
        ct = part.get_content_type()
        try:
            if ct == "text/plain":
                parts.append(part.get_content())
            elif ct == "text/html":
                html = part.get_content()
                clean = re.sub(r"<[^>]+>", " ", html)
                parts.append(re.sub(r"\s+", " ", clean).strip())
        except Exception:
            pass
    return "\n".join(parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_from_pdf(file_bytes)
    elif ext in ("png", "jpg", "jpeg", "bmp", "tiff"):
        return extract_from_image(file_bytes)
    elif ext == "docx":
        return extract_from_docx(file_bytes)
    elif ext == "eml":
        return extract_from_email(file_bytes)
    else:
        return file_bytes.decode("utf-8", errors="ignore")
